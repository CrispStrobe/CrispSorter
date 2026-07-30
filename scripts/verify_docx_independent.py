#!/usr/bin/env python
"""Independent verification of the `crispsorter docx` verbs (P30).

The Rust tests assert against crisp-docx's own reader — the same code that
wrote the file. These assertions come from tools that share nothing with it:
python-docx (a separate OOXML implementation) and plain zipfile/XML reads.
Fixtures are authored here rather than by the code under test.

This harness also exists because the DOCX verbs were, until 2026-07-30,
unreachable from the command line at all: `cli::SUBCOMMANDS` did not list
`docx`, so `crispsorter docx …` fell through to the GUI. Nothing short of
actually running the binary would have caught that — the Rust tests call
the command functions directly, and `--help` renders fine either way.

Usage:
    CRISP_CLI=path/to/crispsorter CRISP_WORK=/tmp/docxverify \\
        python scripts/verify_docx_independent.py
"""
import json
import os
import subprocess
import sys
import zipfile
from pathlib import Path

CLI = os.environ.get("CRISP_CLI", "target/debug/crispsorter")
WORK = Path(os.environ.get("CRISP_WORK", "/tmp/crispsorter-docx-verify"))
WORK.mkdir(parents=True, exist_ok=True)

results = []


def check(name, ok, detail=""):
    results.append((name, bool(ok), detail))
    print(f"{'PASS' if ok else 'FAIL'}  {name}" + (f"  — {detail}" if detail else ""))
    return bool(ok)


class Hung:
    """Stand-in for a call that never returned."""
    returncode = -1
    stdout = ""
    stderr = "the process did not exit — a verb missing from cli::SUBCOMMANDS " \
             "launches the GUI instead of running"


def run(args, timeout=90):
    # Always bounded: an unrecognised verb opens the GUI event loop and waits
    # forever. Without a timeout this harness hangs instead of reporting.
    try:
        return subprocess.run([CLI] + args, capture_output=True, text=True, timeout=timeout)
    except subprocess.TimeoutExpired:
        return Hung()


def ran(name, p):
    if p.returncode != 0:
        msg = (p.stderr or p.stdout).strip()
        check(f"{name} ran", False, msg.splitlines()[-1][:200] if msg else f"exit {p.returncode}")
        return False
    return True


# ── Fixtures, authored the way Word would not, but a reader must accept ──
CT = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>'
ROOT_RELS = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>'
DOC_RELS = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'


def docx(body: str, path: Path) -> Path:
    doc = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"><w:body>'
        + body + "</w:body></w:document>"
    ).encode()
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", CT)
        z.writestr("_rels/.rels", ROOT_RELS)
        z.writestr("word/document.xml", doc)
        z.writestr("word/_rels/document.xml.rels", DOC_RELS)
    return path


def para(text, bold=False, size=None):
    rpr = ""
    if bold or size:
        rpr = "<w:rPr>" + ("<w:b/>" if bold else "") + (f'<w:sz w:val="{size}"/>' if size else "") + "</w:rPr>"
    return f"<w:p>{('<w:pPr>' + rpr + '</w:pPr>') if rpr else ''}<w:r>{rpr}<w:t xml:space=\"preserve\">{text}</w:t></w:r></w:p>"


def python_docx_paragraphs(path):
    """Paragraph texts as a *different* OOXML implementation sees them."""
    import docx as pydocx
    return [p.text for p in pydocx.Document(str(path)).paragraphs]


def part(path, name):
    with zipfile.ZipFile(path) as z:
        return z.read(name).decode("utf-8", "replace")


def part_names(path):
    with zipfile.ZipFile(path) as z:
        return z.namelist()


def main():
    if not Path(CLI).exists():
        print(f"CLI not found: {CLI}", file=sys.stderr)
        return 2

    # A document with an inferable outline and straight quotes.
    src = docx(
        para("Chapter One", bold=True, size=48)
        + "".join(para(f'He said "hello" number {i} to them.', size=22) for i in range(3))
        + para("Section A", bold=True, size=32)
        + "".join(para(f"ordinary prose line {i}", size=22) for i in range(3)),
        WORK / "src.docx",
    )
    check("fixture opens in python-docx", len(python_docx_paragraphs(src)) == 8,
          str(len(python_docx_paragraphs(src))))

    print("\n-- docx check --")
    p = run(["--format", "json", "docx", "check", str(src)])
    if ran("docx check", p):
        r = json.loads(p.stdout)
        check("check: a sound package is valid", r["valid"], str(r["issues"])[:120])
        check("check: names the checks it ran", len(r["ok"]) >= 5, str(len(r["ok"])))

    print("\n-- docx info --")
    p = run(["--format", "json", "docx", "info", str(src)])
    if ran("docx info", p):
        b = json.loads(p.stdout)
        check("info: reports a style count", "style_count" in b, str(b)[:80])
        check("info: reports sections", isinstance(b.get("sections"), list))

    print("\n-- docx headings --")
    p = run(["--format", "json", "docx", "headings", str(src)])
    if ran("docx headings", p):
        hs = json.loads(p.stdout)
        levels = {h["text"]: h["level"] for h in hs}
        check("headings: found both bold paragraphs", set(levels) == {"Chapter One", "Section A"},
              str(levels))
        check("headings: the larger one outranks the smaller",
              levels.get("Chapter One", 9) < levels.get("Section A", 0), str(levels))
        # The body prose must NOT be promoted to a heading.
        check("headings: body prose is not a heading",
              not any("ordinary prose" in h["text"] for h in hs))

    print("\n-- docx normalize-quotes --")
    de = WORK / "de.docx"
    if ran("docx normalize-quotes", run(["docx", "normalize-quotes", str(src),
                                         "--style", "german", "--out", str(de)])):
        texts = python_docx_paragraphs(de)
        joined = "\n".join(texts)
        check("quotes: German opener present (python-docx)", "„" in joined, joined[:60])
        check("quotes: straight quotes gone (python-docx)", '"hello"' not in joined)
        check("quotes: the words survived", "number 1 to them" in joined)

    print("\n-- docx strip-rsids --")
    rsid_src = docx(
        '<w:p w14:paraId="A1B2" w:rsidR="00112233" w:rsidRPr="DEADBEEF">'
        '<w:r w:rsidR="11223344"><w:t>tracked text</w:t></w:r></w:p>',
        WORK / "rsid.docx",
    )
    clean = WORK / "clean.docx"
    if ran("docx strip-rsids", run(["docx", "strip-rsids", str(rsid_src), "--out", str(clean)])):
        xml = part(clean, "word/document.xml")
        check("rsids: no rsid attribute remains (raw XML)",
              "w:rsidR" not in xml and "w:rsidRPr" not in xml, xml[:120])
        check("rsids: no paraId remains (raw XML)", "w14:paraId" not in xml)
        check("rsids: text survived (python-docx)",
              "tracked text" in "\n".join(python_docx_paragraphs(clean)))

    print("\n-- docx inject-footnotes --")
    marked = docx(para("opener.[1] middle.[2] end."), WORK / "marked.docx")
    noted = WORK / "noted.docx"
    p = run(["docx", "inject-footnotes", str(marked),
             "--note", "1=first note", "--note", "2=second note", "--out", str(noted)])
    if ran("docx inject-footnotes", p):
        names = part_names(noted)
        check("footnotes: a footnotes part was created",
              any("footnotes.xml" in n for n in names), str(names))
        fx = part(noted, "word/footnotes.xml")
        check("footnotes: note text is in the part", "first note" in fx and "second note" in fx)
        doc_xml = part(noted, "word/document.xml")
        check("footnotes: references were inserted (raw XML)",
              doc_xml.count("footnoteReference") == 2, str(doc_xml.count("footnoteReference")))
        check("footnotes: literal markers are gone (raw XML)",
              "[1]" not in doc_xml and "[2]" not in doc_xml)
        body = "\n".join(python_docx_paragraphs(noted))
        check("footnotes: surrounding text intact (python-docx)",
              "opener." in body and "end." in body, body[:80])
        # A duplicate marker must be refused, not silently overwritten.
        dup = run(["docx", "inject-footnotes", str(marked), "--note", "1=a",
                   "--note", "1=b", "--out", str(WORK / "dup.docx")])
        check("footnotes: a duplicated marker is refused", dup.returncode != 0,
              (dup.stderr or "").strip()[:80])

    print("\n-- docx convert-notes --")
    en = WORK / "endnotes.docx"
    if ran("docx convert-notes", run(["docx", "convert-notes", str(noted),
                                       "--to", "endnotes", "--out", str(en)])):
        names = part_names(en)
        check("notes: an endnotes part exists", any("endnotes.xml" in n for n in names), str(names))
        doc_xml = part(en, "word/document.xml")
        check("notes: body now references endnotes (raw XML)",
              "endnoteReference" in doc_xml and "footnoteReference" not in doc_xml)
        bad = run(["docx", "convert-notes", str(noted), "--to", "sidenotes",
                   "--out", str(WORK / "no.docx")])
        check("notes: an unknown kind is refused", bad.returncode != 0)

    print("\n-- docx restyle --")
    blueprint = docx(
        para("template body to be replaced")
        + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>',
        WORK / "blueprint.docx",
    )
    restyled = WORK / "restyled.docx"
    if ran("docx restyle", run(["docx", "restyle", "--source", str(src),
                                 "--blueprint", str(blueprint), "--out", str(restyled)])):
        texts = "\n".join(python_docx_paragraphs(restyled))
        check("restyle: source content is present (python-docx)", "Chapter One" in texts, texts[:80])
        check("restyle: template body is gone (python-docx)",
              "template body to be replaced" not in texts)
        # Geometry must come from the blueprint. python-docx reports EMU.
        import docx as pydocx
        sec = pydocx.Document(str(restyled)).sections[0]
        check("restyle: blueprint page size kept (python-docx)",
              sec.page_width is not None and abs(sec.page_width.inches - 8.5) < 0.01,
              f"{sec.page_width.inches if sec.page_width else None} in")
        p2 = run(["--format", "json", "docx", "check", str(restyled)])
        if p2.returncode == 0:
            check("restyle: output still validates", json.loads(p2.stdout)["valid"],
                  str(json.loads(p2.stdout)["issues"])[:120])

    print()
    passed = sum(1 for _, ok, _ in results if ok)
    print(f"{passed}/{len(results)} independent checks passed")
    failed = [(n, d) for n, ok, d in results if not ok]
    if failed:
        print("\nFAILED:")
        for n, d in failed:
            print(f"  - {n}" + (f"  ({d})" if d else ""))
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
